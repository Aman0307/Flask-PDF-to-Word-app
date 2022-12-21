from flask import Flask, request, render_template, send_file
from werkzeug.utils import secure_filename
from docx import Document
import PyPDF2

app = Flask(__name__)

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/convert", methods=["POST"])
def convert():
    # Check if the POST request has the file part
    if "file" not in request.files:
        return "No file part"
    file = request.files["file"]
    # If the user did not select a file, return an error message
    if file.filename == "":
        return "No selected file"
    # Save the file to the server and secure the filename
    filename = secure_filename(file.filename)
    file.save(filename)
    # Convert the PDF to a Word document
    document = Document()
    with open(filename, "rb") as fp:
        # Create a PDF object
        pdf = PyPDF2.PdfFileReader(fp)
        # Create a Word document
        document = Document()
        # Iterate over all pages in the PDF
        for page in range(pdf.getNumPages()):
            # Extract the text from the page and add it to the Word document
            text = pdf.getPage(page).extractText()
            document.add_paragraph(text)
    # Save the Word document
    document.save("converted.docx")
    # Send the converted file as an attachment to the client
    return send_file('converted.docx', attachment_filename='converted.docx', as_attachment=True)



if __name__ == "__main__":
    app.run()
